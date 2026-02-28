"""
Document Format Conversion Utilities
Handles conversion to Markdown, HTML, DOCX while preserving formatting
"""

import re
from typing import List, Dict, Any
from io import BytesIO
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import markdown
import base64
from PIL import Image


class DocumentConverter:
    """Handles conversion of OCR results to various document formats"""

    def __init__(self):
        self.page_separator = '<--- Page Split --->'

    def to_markdown(self, pages_content: List[Dict[str, Any]], include_images: bool = True) -> str:
        """
        Convert OCR results to Markdown format

        Args:
            pages_content: List of page dictionaries with text and metadata
            include_images: Whether to include image references

        Returns:
            Markdown formatted string
        """
        md_content = []

        for idx, page in enumerate(pages_content):
            # Add page header
            md_content.append(f"# Page {idx + 1}\n")

            text = page.get('text', '')

            # Process and clean the text
            if include_images and 'images' in page:
                # Replace image placeholders with actual markdown image syntax
                for img_idx, img_data in enumerate(page.get('images', [])):
                    placeholder = f"[IMAGE_{img_idx}]"
                    img_ref = f"![Image {img_idx + 1}](data:image/jpeg;base64,{img_data})"
                    text = text.replace(placeholder, img_ref)

            md_content.append(text)
            md_content.append("\n\n---\n\n")  # Page separator

        return "\n".join(md_content)

    def to_html(self, pages_content: List[Dict[str, Any]], include_images: bool = True) -> str:
        """
        Convert OCR results to HTML format

        Args:
            pages_content: List of page dictionaries with text and metadata
            include_images: Whether to include images

        Returns:
            HTML formatted string
        """
        html_parts = []

        # HTML header
        html_parts.append("""
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>OCR Results</title>
    <style>
        body {
            font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
            max-width: 900px;
            margin: 40px auto;
            padding: 20px;
            line-height: 1.6;
            background-color: #f5f5f5;
        }
        .page {
            background: white;
            padding: 40px;
            margin-bottom: 30px;
            box-shadow: 0 2px 8px rgba(0,0,0,0.1);
            border-radius: 8px;
        }
        .page-header {
            color: #333;
            border-bottom: 2px solid #4CAF50;
            padding-bottom: 10px;
            margin-bottom: 20px;
        }
        table {
            border-collapse: collapse;
            width: 100%;
            margin: 20px 0;
        }
        th, td {
            border: 1px solid #ddd;
            padding: 12px;
            text-align: left;
        }
        th {
            background-color: #4CAF50;
            color: white;
        }
        tr:nth-child(even) {
            background-color: #f9f9f9;
        }
        img {
            max-width: 100%;
            height: auto;
            margin: 15px 0;
            border-radius: 4px;
        }
        code {
            background-color: #f4f4f4;
            padding: 2px 6px;
            border-radius: 3px;
            font-family: 'Courier New', monospace;
        }
        pre {
            background-color: #f4f4f4;
            padding: 15px;
            border-radius: 5px;
            overflow-x: auto;
        }
    </style>
</head>
<body>
    <h1>DeepSeek OCR Results</h1>
""")

        # Process each page
        for idx, page in enumerate(pages_content):
            html_parts.append(f'    <div class="page">')
            html_parts.append(f'        <h2 class="page-header">Page {idx + 1}</h2>')

            text = page.get('text', '')

            # Handle images if present
            if include_images and 'images' in page:
                for img_idx, img_data in enumerate(page.get('images', [])):
                    placeholder = f"[IMAGE_{img_idx}]"
                    img_tag = f'<img src="data:image/jpeg;base64,{img_data}" alt="Image {img_idx + 1}" />'
                    text = text.replace(placeholder, img_tag)

            # Convert markdown to HTML if the text appears to be markdown
            if self._is_markdown(text):
                html_content = markdown.markdown(text, extensions=['tables', 'fenced_code'])
            else:
                # Otherwise, preserve the HTML or wrap in paragraph
                html_content = text if '<' in text else f'<p>{text.replace(chr(10), "<br>")}</p>'

            html_parts.append(f'        {html_content}')
            html_parts.append('    </div>')

        # HTML footer
        html_parts.append("""
</body>
</html>
""")

        return "\n".join(html_parts)

    def to_docx(self, pages_content: List[Dict[str, Any]], include_images: bool = True) -> BytesIO:
        """
        Convert OCR results to DOCX format

        Args:
            pages_content: List of page dictionaries with text and metadata
            include_images: Whether to include images

        Returns:
            BytesIO object containing the DOCX file
        """
        doc = Document()

        # Set default font
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)

        # Add title
        title = doc.add_heading('DeepSeek OCR Results', 0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Process each page
        for idx, page in enumerate(pages_content):
            # Add page heading
            page_heading = doc.add_heading(f'Page {idx + 1}', level=1)
            page_heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

            text = page.get('text', '')

            # Handle images
            if include_images and 'images' in page:
                for img_idx, img_data in enumerate(page.get('images', [])):
                    placeholder = f"[IMAGE_{img_idx}]"

                    # Add image to document
                    try:
                        img_bytes = base64.b64decode(img_data)
                        img_stream = BytesIO(img_bytes)
                        doc.add_picture(img_stream, width=Inches(5))
                        text = text.replace(placeholder, '')
                    except Exception as e:
                        print(f"Error adding image to DOCX: {e}")

            # Process text content
            self._add_formatted_text_to_doc(doc, text)

            # Add page break (except for last page)
            if idx < len(pages_content) - 1:
                doc.add_page_break()

        # Save to BytesIO
        docx_buffer = BytesIO()
        doc.save(docx_buffer)
        docx_buffer.seek(0)

        return docx_buffer

    def _is_markdown(self, text: str) -> bool:
        """Check if text appears to be markdown formatted"""
        markdown_patterns = [
            r'^#+\s',  # Headers
            r'\*\*.*\*\*',  # Bold
            r'\*.*\*',  # Italic
            r'^\*\s',  # Lists
            r'^\d+\.\s',  # Numbered lists
            r'\[.*\]\(.*\)',  # Links
            r'```',  # Code blocks
        ]

        for pattern in markdown_patterns:
            if re.search(pattern, text, re.MULTILINE):
                return True
        return False

    def _add_formatted_text_to_doc(self, doc: Document, text: str):
        """
        Add formatted text to document, preserving structure

        Args:
            doc: Document object
            text: Text to add
        """
        # Split into paragraphs
        paragraphs = text.split('\n\n')

        for para in paragraphs:
            if not para.strip():
                continue

            # Check for headers
            if para.startswith('# '):
                doc.add_heading(para.replace('# ', ''), level=1)
            elif para.startswith('## '):
                doc.add_heading(para.replace('## ', ''), level=2)
            elif para.startswith('### '):
                doc.add_heading(para.replace('### ', ''), level=3)
            # Check for tables (simple detection)
            elif '|' in para and para.count('|') > 2:
                self._add_table_to_doc(doc, para)
            # Check for code blocks
            elif para.startswith('```'):
                code_text = para.strip('```').strip()
                p = doc.add_paragraph()
                run = p.add_run(code_text)
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
            else:
                # Regular paragraph
                doc.add_paragraph(para.strip())

    def _add_table_to_doc(self, doc: Document, table_text: str):
        """
        Add a table to the document from markdown-style table text

        Args:
            doc: Document object
            table_text: Table in markdown format
        """
        rows = [row.strip() for row in table_text.split('\n') if row.strip()]

        # Filter out separator rows
        data_rows = [row for row in rows if not re.match(r'^[\|\s\-:]+$', row)]

        if not data_rows:
            return

        # Parse table data
        table_data = []
        for row in data_rows:
            cells = [cell.strip() for cell in row.split('|')]
            cells = [c for c in cells if c]  # Remove empty cells
            if cells:
                table_data.append(cells)

        if not table_data:
            return

        # Create table
        max_cols = max(len(row) for row in table_data)
        table = doc.add_table(rows=len(table_data), cols=max_cols)
        table.style = 'Light Grid Accent 1'

        # Populate table
        for i, row_data in enumerate(table_data):
            row = table.rows[i]
            for j, cell_text in enumerate(row_data):
                if j < len(row.cells):
                    row.cells[j].text = cell_text

                    # Make header row bold
                    if i == 0:
                        for paragraph in row.cells[j].paragraphs:
                            for run in paragraph.runs:
                                run.font.bold = True
